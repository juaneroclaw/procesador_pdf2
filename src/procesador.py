import os
import re
from pdf2image import convert_from_path
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
import pdfplumber

class ProcesadorPDFOptimizado:
    def __init__(self, ruta_pdf, directorio_salida='salida'):
        self.ruta_pdf = ruta_pdf
        self.directorio_salida = directorio_salida
        os.makedirs(self.directorio_salida, exist_ok=True)

    def dividir_pdf(self, paginas_por_archivo=10):
        try:
            lector_pdf = PdfReader(self.ruta_pdf)
            total_paginas = len(lector_pdf.pages)
            
            pdfs_generados = []
            for inicio in range(0, total_paginas, paginas_por_archivo):
                escritor = PdfWriter()
                
                fin = min(inicio + paginas_por_archivo, total_paginas)
                for i in range(inicio, fin):
                    escritor.add_page(lector_pdf.pages[i])
                
                nombre_archivo = os.path.join(
                    self.directorio_salida, 
                    f'parte_{inicio//paginas_por_archivo + 1}.pdf'
                )
                
                with open(nombre_archivo, 'wb') as archivo_salida:
                    escritor.write(archivo_salida)
                
                pdfs_generados.append(nombre_archivo)
            
            return pdfs_generados
        
        except Exception as e:
            print(f"Error al dividir PDF: {e}")
            return []

    def convertir_a_imagenes(self, ruta_pdf, parte_numero, formato='webp', resolucion=100):
        try:
            imagenes = convert_from_path(
                ruta_pdf, 
                dpi=resolucion,
                thread_count=4
            )
            
            rutas_imagenes = []
            for i, imagen in enumerate(imagenes):
                ruta_imagen = os.path.join(
                    self.directorio_salida, 
                    f'parte_{parte_numero}_pagina_{i+1}.{formato}'
                )
                
                imagen.save(
                    ruta_imagen, 
                    formato.upper(), 
                    optimize=True,
                    quality=50
                )
                
                rutas_imagenes.append(ruta_imagen)
            
            return rutas_imagenes
        
        except Exception as e:
            print(f"Error al convertir imágenes: {e}")
            return []

    def extraer_texto_a_word(self, ruta_pdf, parte_numero):
        try:
            # Método 1: Intentar con pdfplumber primero
            try:
                with pdfplumber.open(ruta_pdf) as pdf:
                    documento = Document()
                    documento.add_heading(f'Texto - Parte {parte_numero}', level=1)
                    
                    for pagina in pdf.pages:
                        texto = pagina.extract_text()
                        if texto:
                            texto_limpio = self.limpiar_texto(texto)
                            documento.add_paragraph(texto_limpio)
            
            # Si pdfplumber falla, usar PyPDF2 como respaldo
            except Exception as e:
                print(f"Error con pdfplumber: {e}. Usando PyPDF2 como respaldo.")
                lector = PdfReader(ruta_pdf)
                documento = Document()
                documento.add_heading(f'Texto - Parte {parte_numero}', level=1)
                
                for pagina in lector.pages:
                    texto = pagina.extract_text()
                    if texto:
                        texto_limpio = self.limpiar_texto(texto)
                        documento.add_paragraph(texto_limpio)
            
            # Guardar documento
            ruta_word = os.path.join(
                self.directorio_salida, 
                f'texto_parte_{parte_numero}.docx'
            )
            documento.save(ruta_word)
            return ruta_word
        
        except Exception as e:
            print(f"Error al extraer texto: {e}")
            return None

    def limpiar_texto(self, texto):
        """
        Método para limpiar y mejorar la extracción de texto
        """
        try:
            # Eliminar saltos de línea múltiples
            texto = re.sub(r'\n+', '\n', texto)
            
            # Eliminar guiones que cortan palabras
            texto = re.sub(r'(\w+)-\n(\w+)', r'\1\2', texto)
            
            # Eliminar espacios extra
            texto = re.sub(r'\s+', ' ', texto).strip()
            
            # Opcional: Corregir formato de párrafos
            texto = re.sub(r'([.!?])\s*([A-Z])', r'\1\n\2', texto)
            
            return texto
        
        except Exception as e:
            print(f"Error al limpiar texto: {e}")
            return texto
